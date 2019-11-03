import requests                             # for making HTML requests
from bs4 import BeautifulSoup               # BS for parsing HTML
import re                                   # regex
from datetime import datetime as dttm         # datetime for logging
import os
import json
import docx
from ast import literal_eval
import time
from docx.shared import Pt
import docxToPDF
import logging


## GLOBALS ##
# Log file path
LOG_PATH = '~/gktoday/log/'
LOG_PATH = os.path.expanduser(LOG_PATH)
# main quiz page url
main_page = "https://www.gktoday.com/category/gk-current-affairs-quiz-questions-answers/page/{page}/"
quiz_data_url = "https://www.gktoday.com/wp-content/plugins/wp-quiz-basic/ajax_requests.php"

# headers for requests
req_head = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/77.0.3865.120 Safari/537.36'}

# PATHS for saving docx,pdf files
quiz_folder = "~/gktoday/res/quiz/"
docx_folder = "docx/"
pdf_folder = "pdf/"
# TO expand the "~/ path"
quiz_folder = os.path.expanduser(quiz_folder)
daily_quiz_metadata = {}
all_files_metadata = {}


def make_directories():
    # to create the folder for docx and pdf files if they doesn't exist already
    try:
        os.makedirs(quiz_folder + "docx/", exist_ok=True)
        os.makedirs(quiz_folder + "pdf/", exist_ok=True)
        os.makedirs(LOG_PATH, exist_ok=True)
        return 1
    except:
        print("Error while making docx, pdf directories.")
        return 0


def convert_to_pdf(docx_path):
    try:
        docxToPDF.convert_to(quiz_folder + "pdf/", docx_path)

    except:
        print("Error occured in convert_to_pdf() at", dttm.now())
        logging.exception('Error In convert_to_pdf()')
        raise

# string cleanup / remove errors caused by ' or \" or \'
def l_eval(string):
    try:
        return literal_eval("'%s'" %"".join(string.splitlines()))
    except SyntaxError:
        string = string.replace("'",r"\'")
        return literal_eval("'%s'" %"".join(string.splitlines()))


# write data to the docx file
def write_to_docx(path, file_name, ques_data):
    try:
        os.makedirs(path, exist_ok=True)
        try:
            doc = docx.Document(path + file_name)
        except docx.opc.exceptions.PackageNotFoundError:        # file doesn't exist
            doc = docx.Document()
        except FileNotFoundError:
            doc = docx.Document()

        dict_map = {'1': '(a)', '2': '(b)', '3': '(c)', '4': '(d)'}

        p = doc.add_paragraph('')
        if(ques_data['date']):
            date_run = p.add_run(ques_data['date'])             # Date of the current Quiz
            font = date_run.font
            font.bold = True
            font.underline = True
            font.size = Pt(16)
            font.color.rgb = docx.shared.RGBColor(0xFF, 0x0, 0x0)
            p.add_run('\n\n')
        p.add_run('Question ' + str(ques_data['question_counter']) + ': ').bold = True
        p.add_run(ques_data['qu'] + '\n' + '(a) ' + ques_data['opt1'] + '\n' + '(b) ' + ques_data['opt2'] + '\n' + '(c) ' + ques_data['opt3'] + '\n' + '(d) ' + ques_data['opt4'] + '\n')
        p.add_run( 'Answer: ' + dict_map[ques_data['ans']] + '\n' + 'Answer Explanation: ').bold = True
        p.add_run(ques_data['ans_exp']).italic = True
        p.add_run('\n\n')

        doc.save(path + file_name)      # save the docx file

    except:
        print("Error occured in write_to_docx() at", dttm.now())
        logging.exception('Error In write_to_docx()')
        raise

# write data to the file
def write_data(path, data):
    try:
        os.makedirs(quiz_folder, exist_ok=True)
        f = open(path, "w")
        f.write(data)
        f.close()

    except:
        print("Error occured in write_data() at", dttm.now())
        logging.exception('Error In write_data()')
        raise


# initialize the daily_quiz_metadata dictionary
def init_metadata():
    try:
        make_directories()
        global daily_quiz_metadata
        global all_files_metadata
        try:
            f = open(quiz_folder + "daily_quiz_metadata.txt", "r")
            g = open(quiz_folder + "all_files_metadata.txt", "r")
            all_files_metadata = json.loads(g.read())
            daily_quiz_metadata = json.loads(f.read())
            f.close()
            g.close()
        except FileNotFoundError:
            pass

    except:
        print("Error occured in open_daily_quiz_metadata() at", dttm.now())
        logging.exception('Error In init_metadata()')
        raise

# get year month and data from the string
def get_year_month_date(strng):
    try:
        year_re = r'\d{4}'
        month_re = r'^[a-zA-Z]*'
        date_re = r'[0-9].*,'
        yr = re.findall(year_re,strng)[0]
        mth = re.findall(month_re,strng)[0]
        try:
            dt = re.findall(date_re, strng)[0]
        except IndexError:
            dt = strng.replace(yr, "")
            dt = dt.replace(mth, "")
            dt = dt.lstrip()
            dt = dt.rstrip()
        dt = dt.replace(",","")

        return yr,mth,dt

    except:
        print("Error occured in get_year_month_date() at", dttm.now())
        logging.exception('Error In get_year_month_date()')
        raise



# fetch questions' ID from the link
def fetch_ques_id(quiz_link):

    try:
        # request page using `requests`
        req_page = requests.get(url=quiz_link, headers=req_head)
        html = req_page.text

        # parse page using beautiful soup
        soup = BeautifulSoup(html, 'html.parser')

        # get all questions' ID from the current page
        ques_ids = soup.find('form').get('data-quesids')

        return ques_ids
    except:
        print("Error occured in fetch_ques_id() at", dttm.now())
        logging.exception('Error In fetch_ques_id()')
        raise

# fetch quiz data using the ajax url with parameters `params`
def fetch_quiz_data(params):
    try:
        r = requests.post(url=quiz_data_url, data=params, headers=req_head)
        return r
        
    except:
        print("Error occured in fetch_quiz_data() at", dttm.now())
        logging.exception('Error In fetch_quiz_data()')
        raise

def extract_quiz_from_json(ques, date_of_article, yr, mth):
    question_counter = 1
    try:
        for que in ques:

            docx_data = {}
            docx_data['qu'] = l_eval(que['question'])
            docx_data['opt1'] = l_eval(que['option1'])
            docx_data['opt2'] = l_eval(que['option2'])
            docx_data['opt3'] = l_eval(que['option3'])
            docx_data['opt4'] = l_eval(que['option4'])
            docx_data['ans_exp'] = l_eval(que['ans_explanation'])
            try:
                docx_data['ans'] = l_eval(que['answer'])
            except KeyError:
                raise
            docx_data['question_counter'] = question_counter
            if(question_counter == 1):
                docx_data['date'] = date_of_article
            else:
                docx_data['date'] = False

            docx_path = quiz_folder + docx_folder
            docx_file_name = yr + "-" + mth + ".docx"
            write_to_docx(docx_path, docx_file_name, docx_data)

            question_counter += 1           # increment question counter
        
    except :
        print("Error occured in extract_quiz_from_json() at", dttm.now())
        logging.exception('Error In extract_quiz_from_json()')
        raise

    return 0

# convert new docx files to pdf
def convert_new_docx_to_pdf():
    for _year,_year_val in all_files_metadata.items():
        for _month,_month_val in _year_val.items():
            if(not (_month_val['to_pdf'])):
                docx_file_name = _month_val['file_name'] + '.docx'
                docx_path = quiz_folder + "docx/" + docx_file_name
                convert_to_pdf(docx_path)
                all_files_metadata[_year][_month]['to_pdf'] = 'True'
                all_files_md_path = quiz_folder + "all_files_metadata.txt"
                write_data(all_files_md_path, json.dumps(all_files_metadata))

# extract all the article (date and urls) from the main_page
def fetch_articles(year):
    yr = year[-1]
    pg = 1
    try:
        while(True):

            MAIN_PAGE = main_page.format(page=pg)

            # request page using `requests`
            req_page = requests.get(url=MAIN_PAGE, headers= req_head)
            html = req_page.text

            # parse page using beautiful soup
            soup = BeautifulSoup(html, 'html.parser')

            # get all articles from the current page
            all_articles = soup.find_all('article')
            for article in all_articles:

                all_h1 = article.find_all('h1')
                date_of_article = all_h1[0].find('a').text
                link = all_h1[0].find('a').get('href')
                date_of_article = date_of_article.replace('GK & Current Affairs Quiz: ', "")

                # get year, month, date of the current article
                yr, mth, dt = get_year_month_date(date_of_article)

                # skip when the current year is not the specified year
                if(not (int(yr) in year)):
                    continue
                # skip weekly quizes
                if(mth == "Weekly"):
                    continue
                to_write = False
                if(not (yr in all_files_metadata)):
                    all_files_metadata[yr] = {}
                    all_files_metadata[yr][mth] = {}
                    to_write = True

                elif(not (mth in all_files_metadata[yr])):
                    all_files_metadata[yr][mth] = {}
                    to_write = True


                to_write_daily_md = False
                if(not (yr in daily_quiz_metadata)):        # if year is not present in the metadata
                    daily_quiz_metadata[yr] = {}
                    daily_quiz_metadata[yr][mth] = {}
                    daily_quiz_metadata[yr][mth][dt] = {}
                    to_write_daily_md = True

                elif(not (mth in daily_quiz_metadata[yr])):  # if year is present but month isn't
                    daily_quiz_metadata[yr][mth] = {}
                    daily_quiz_metadata[yr][mth][dt] = {}
                    to_write_daily_md = True

                elif (not (dt in daily_quiz_metadata[yr][mth])): # if year & month are present but date isn't
                    daily_quiz_metadata[yr][mth][dt] = {}
                    to_write_daily_md = True

                # write current time to the year>month>date dictionary to track when the quiz is fetched for the first time from the server
                if(to_write_daily_md):
                    ques_ids = fetch_ques_id(link)
                    daily_quiz_metadata[yr][mth][dt]['time'] = dttm.now().strftime("%d/%b/%Y, %H:%M:%S")
                    daily_quiz_metadata[yr][mth][dt]['url'] = link
                    daily_quiz_metadata[yr][mth][dt]['ques_ids'] = ques_ids

                    # write daily_quiz_metadata to .txt file
                    daily_quiz_txt_path = quiz_folder + "daily_quiz_metadata.txt"
                    all_files_md_path = quiz_folder + "all_files_metadata.txt"

                    if(to_write):
                        file_nm = yr + "-" + mth
                        all_files_metadata[yr][mth]['file_name'] = file_nm
                        all_files_metadata[yr][mth]['file_path'] = quiz_folder + pdf_folder + file_nm + ".pdf"
                    all_files_metadata[yr][mth]['last_updated'] = dttm.now().strftime("%d/%b/%Y, %H:%M:%S")
                    all_files_metadata[yr][mth]['to_pdf'] = False
                    write_data(all_files_md_path, json.dumps(all_files_metadata))

                    params = {'quesids': ques_ids,'subaction': 'submitquiz','testtype': 'daily', 'showans': '1'}

                    r = fetch_quiz_data(params)

                    json_res = json.loads(r.content)
                    if(json_res['status'] == True):

                        # we successfully received the quizdata for this date
                        ques = json_res['questions']
                        
                        x = extract_quiz_from_json(ques, date_of_article, yr, mth)
                        print(dt, mth, yr)
                        if(not x):
                            write_data(daily_quiz_txt_path, json.dumps(daily_quiz_metadata))
                else:
                    print("skip", dt, mth, yr)
            if(not (int(yr) in year)):
                break
            pg += 1
        
    except :
        print("Error occured in fetch_articles() at", dttm.now())
        logging.exception('Error In fetch_articles()')
        raise



## __MAIN__

if(__name__ == "__main__"):
    try:
        # initialize metadata
        init_metadata()
        logging.basicConfig(filename=LOG_PATH+"scraping.log", level=logging.DEBUG)
        # fetch and write quiz data to docx file
        y = list(range(2019,2020))
        fetch_articles(y)

        # Convert new docx files to pdf
        convert_new_docx_to_pdf()
        
        print("scraping successfull!")

    except:
        print("Error occured in __main__ at", dttm.now())
        logging.exception('Error In __main__')
        raise
