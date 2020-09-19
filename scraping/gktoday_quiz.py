import json
import logging
import os
from datetime import datetime as dttm

import docx
import requests
from bs4 import BeautifulSoup
from docx.shared import Pt

from scraping.gk_utility import docxToPDF
from scraping.gk_utility import gk_utils
from scraping import QUIZ_LOG_PATH, product_url, docx_folder, pdf_folder, main_page_url, main_page_json
from scraping import quiz_folder, fortnight_folder, monthly_folder, month_map
from scraping.gk_utility import req_head

# GLOBALS
# Log file path
LOG_PATH = QUIZ_LOG_PATH
LOG_PATH = os.path.expanduser(LOG_PATH)
# TO expand the "~/ path"
quiz_folder = os.path.expanduser(quiz_folder)
all_files_metadata = {}


def make_directories():
    # to create the folder for docx and pdf files if they doesn't exist already
    logging.info("making directories")
    os.makedirs(quiz_folder + "docx/", exist_ok=True)
    os.makedirs(quiz_folder + "pdf/monthly/", exist_ok=True)
    os.makedirs(quiz_folder + "pdf/fortnight/", exist_ok=True)
    os.makedirs(quiz_folder + "pdf/upsc_prelims/", exist_ok=True)
    os.makedirs(LOG_PATH, exist_ok=True)
    return 1


def convert_to_pdf(docx_path, quiz_type):
    logging.debug("inside convert_new_docx_to_pdf()")
    try:
        docxToPDF.convert_to(quiz_folder + pdf_folder + quiz_type, docx_path)
        name = docx_path.replace(quiz_folder + docx_folder, "")
        logging.info("Successfully converted `" + name + "` to PDF.")
    except docxToPDF.LibreOfficeError:
        logging.exception('Error In convert_to_pdf()')
        raise docxToPDF.LibreOfficeError


# write data to the docx file
def write_to_docx(path, file_name, ques_list):
    os.makedirs(path, exist_ok=True)
    try:
        doc = docx.Document(path + file_name)
    except docx.opc.exceptions.PackageNotFoundError:  # file doesn't exist
        doc = docx.Document()
    except FileNotFoundError:
        doc = docx.Document()

    for ques_data in ques_list:
        p = doc.add_paragraph('')
        if int(ques_data['question_counter']) % 10 == 0:
            advert = p.add_run("@GKTodayQuizBot on Telegram")
            font = advert.font
            font.bold = True
            font.size = Pt(25)
            font.name = 'forte'
            p.paragraph_format.alignment = 1
            font.color.rgb = docx.shared.RGBColor(0x00, 0xB0, 0x50)
            p.add_run('\n\n')
        p = doc.add_paragraph('')
        p.paragraph_format.alignment = 0
        p.add_run('Question ' + str(ques_data['question_counter']) + ': ').bold = True
        p.add_run(
            ques_data['qu'] + '\n' + '(a) ' + ques_data['opt1'] + '\n' + '(b) ' + ques_data[
                'opt2'] + '\n' + '(c) ' +
            ques_data['opt3'] + '\n' + '(d) ' + ques_data['opt4'] + '\n')
        p.add_run('Answer: ' + ques_data['ans'] + '\n' + 'Answer Explanation: ').bold = True
        p.add_run(ques_data['ans_exp']).italic = True

    doc.save(path + file_name)  # save the docx file


# write data to the file
def write_data(path, data):
    os.makedirs(quiz_folder, exist_ok=True)
    f = open(path, "w")
    f.write(data)
    f.close()


# initialize the daily_quiz_metadata dictionary
def init_metadata():
    make_directories()
    global all_files_metadata
    try:
        g = open(quiz_folder + "all_files_metadata.json", "r")
        all_files_metadata = json.loads(g.read())
        g.close()
    except FileNotFoundError:
        logging.info("Couldn't read all_files_metadata.json because it doesn't exist in FS")


# function to create the docx data object which will be used by the write_to_docx function to create the docx file
# INPUT:
#   product_id : product_id of the product which is to be fetched
def create_docx_obj(product_id, path, file_name):
    data = fetch_product(product_id)
    data_list = data.split("<span class=\"quesno\">")[1:]
    ques_list = []
    # traverse all the questions
    try:
        for i in data_list:
            soup = BeautifulSoup(i, "lxml")
            try:
                question_counter, ques = soup.find('p').decode_contents().strip().split('. ', 1)
            except ValueError:
                logging.info("Fallback!! ValueError in soup ques product_id: " + soup.text[0:5] + "Product ID: " + str(
                    product_id) + "\nUsing second method to get the required data")
                question_counter = soup.find('p').decode_contents().strip().split('.')[0]
                ques = soup.find('p').next_sibling.text.strip().replace('<br/>', "")
            options = [x[4:] for x in
                       soup.find('div', {"class": "wp_quiz_question_options"}).decode_contents().strip().split('<br/>')]
            exp = soup.find('div', {"class": "exp"}).decode_contents().strip()
            exp = exp.replace("\r", "").strip()
            ans = soup.find("div", {"type": "A"}).next_sibling.text.strip()[8:]
            docx_data = {
                'qu': ques,
                'opt1': options[0],
                'opt2': options[1],
                'opt3': options[2],
                'opt4': options[3],
                'ans_exp': exp,
                'ans': ans,
                'question_counter': question_counter
            }
            ques_list.append(docx_data)
        write_to_docx(path, file_name, ques_list)
    except ValueError:
        logging.exception("ValueError in create_docx_obj()")
        raise Exception()


# fetch the product from the server
# INPUT::
#   product_id : product_id of the product which is to be fetched
def fetch_product(product_id):
    fetch_url = product_url.format(product_id=product_id)
    res = requests.get(url=fetch_url, headers=req_head)
    if res.status_code != 200:
        # error
        raise Exception("Status code" + str(res.status_code))

    return json.loads(res.text)['content']['rendered']


# fetch all products which haven't been fetched yet
def fetch_all_products(to_pdf=False):
    logging.debug("inside fetch_all_products")

    try:
        for typ1, typ1_val in all_files_metadata.items():
            if typ1 == "current_affairs":
                for typ2, typ2_val in typ1_val.items():
                    if typ2 == "monthly":
                        for _yr, _yr_val in typ2_val.items():
                            for _mth, _mth_val in _yr_val.items():
                                _product_id = _mth_val['product_id']
                                # now fetch the product from the server, fetch_product returns an BS4 object of content
                                docx_file_name = _mth + ", " + _yr + ".docx"
                                docx_path = quiz_folder + docx_folder + monthly_folder
                                if not all_files_metadata[typ1][typ2][_yr][_mth]['fetched']:
                                    create_docx_obj(_product_id, docx_path, docx_file_name)
                                    all_files_metadata[typ1][typ2][_yr][_mth]['fetched'] = True
                                    all_files_metadata[typ1][typ2][_yr][_mth]['last_update'] = str(dttm.now())
                                    write_data(quiz_folder + "all_files_metadata.json", json.dumps(all_files_metadata))
                                else:
                                    logging.debug(docx_file_name + " has already been fetched")
                                docx_path += docx_file_name
                                # convert to pdf only when the parameter is true
                                if to_pdf and not all_files_metadata[typ1][typ2][_yr][_mth]['to_pdf']:
                                    convert_to_pdf(docx_path, "monthly")
                                    all_files_metadata[typ1][typ2][_yr][_mth]['to_pdf'] = True
                                    pdf_path = docx_path.replace('/docx/', '/pdf/').replace('.docx', '.pdf')
                                    all_files_metadata[typ1][typ2][_yr][_mth]['pdf_file_path'] = pdf_path
                                    all_files_metadata[typ1][typ2][_yr][_mth]['last_update'] = str(dttm.now())
                                    write_data(quiz_folder + "all_files_metadata.json", json.dumps(all_files_metadata))
                                elif to_pdf:
                                    logging.debug(docx_file_name + " has already been converted to PDF")

                        # end "monthly"
                    elif typ2 == "fortnight":
                        for _yr, _yr_val in typ2_val.items():
                            for _mth, _mth_val in _yr_val.items():
                                for _dt, _dt_val in _mth_val.items():
                                    _product_id = _dt_val['product_id']
                                    docx_file_name = _dt + " " + _mth + ", " + _yr + ".docx"
                                    docx_path = quiz_folder + docx_folder + fortnight_folder
                                    # now fetch the product from the server, fetch_product returns an BS4 object of
                                    # content
                                    if not all_files_metadata[typ1][typ2][_yr][_mth][_dt]['fetched']:
                                        create_docx_obj(_product_id, docx_path, docx_file_name)
                                        all_files_metadata[typ1][typ2][_yr][_mth][_dt]['fetched'] = True
                                        all_files_metadata[typ1][typ2][_yr][_mth][_dt]['last_update'] = str(dttm.now())
                                        write_data(quiz_folder + "all_files_metadata.json",
                                                   json.dumps(all_files_metadata))
                                    else:
                                        logging.debug(docx_file_name + " has already been fetched")
                                    docx_path += docx_file_name
                                    # convert to pdf only when the parameter is true
                                    if to_pdf and not all_files_metadata[typ1][typ2][_yr][_mth][_dt]['to_pdf']:
                                        convert_to_pdf(docx_path, "fortnight")
                                        all_files_metadata[typ1][typ2][_yr][_mth][_dt]['to_pdf'] = True
                                        pdf_path = docx_path.replace('/docx/', '/pdf/').replace('.docx', '.pdf')
                                        all_files_metadata[typ1][typ2][_yr][_mth][_dt]['pdf_file_path'] = pdf_path
                                        all_files_metadata[typ1][typ2][_yr][_mth][_dt]['last_update'] = str(dttm.now())
                                        write_data(quiz_folder + "all_files_metadata.json",
                                                   json.dumps(all_files_metadata))
                                    elif to_pdf:
                                        logging.debug(docx_file_name + " has already been converted to PDF")

                        # end "fortnight"
                    elif typ2 == "upsc_prelims":
                        for _yr, _yr_val in typ2_val.items():
                            for _mth, _mth_val in _yr_val.items():
                                _product_id = _mth_val['product_id']
                                # now fetch the product from the server, fetch_product returns an BS4 object of
                                # content create_docx_obj(_product_id, quiz_folder+docx_folder+upsc_prelims_folder,
                                # _mth + ", " + _yr + ".docx")
                        # end "upsc_prelims"
    except KeyError:
        logging.exception("Keyerror in fetch_all_products:\n")
        raise Exception()


# convenience function to insert dictionary data
def insert_into_dict(trgt, data):
    if type(data) != dict:
        return
    for key, value in data.items():
        if key in trgt:
            insert_into_dict(trgt[key], value)
        else:
            trgt[key] = value


# This function takes in a list of products and categorize them based on their name
# into monthly, fortnight, upsc_prelims type and returns a flag when there is no new product found in the product list
# INPUT:
#   all_products : a list of products
# RETURN:
#   flag : 1,0 : 1 indicates no new product found
def categorize_product(all_products, classic=False):
    flag = 1
    # get product product_id and product name from the current page
    for product in all_products:
        if classic:
            details = product.find('a', {"class": "button"})
            product_name = details.get("aria-label").encode().split(b"\xe2\x80\x9c")[1].split(b"\xe2\x80\x9d")[
                0].decode()

            product_id = details.get("data-product_id")  # nov 2019 onwards, we have good pdf files (workable)
        else:
            product_id = product['id']
            product_name = BeautifulSoup(product['title']['rendered'], "lxml").text
            # this is required because in case of JSON API call the data we get contains en-dash instead of hyphen
            product_name = product_name.replace(u"\u2013", "-")
            logging.debug(str(product_id) + ": " + product_name)
        logging.debug("Product Name:" + product_name + " & Product ID:" + str(product_id))
        temp = ""
        # this will give monthly quiz PDF
        if product_name.find('MCQs PDF') != -1 and product_name.find('Current Affairs') != -1 and int(
                product_id) > 324341:
            # if there is no new product on the page then break
            logging.debug("Product is of type `MCQs PDF`")
            _yr, _mth, _dt = gk_utils.get_year_month_date(product_name.split("-")[1].replace(" ", ""))
            _mth = month_map[_mth.title()] + ". " + _mth.title()
            try:
                if all_files_metadata["current_affairs"]["monthly"][_yr][_mth]["product_id"] == product_id:
                    logging.debug("Product has already been fetched")
                    continue
            except KeyError:
                # when some key doesn't exist -> that this is a new product on the page
                logging.info("[New Product Monthly] | " + str(product_id) + ": " + product_name)
                flag = 0
                temp = {
                    "current_affairs": {
                        "monthly": {
                            _yr: {
                                _mth: {
                                    "product_id": product_id,
                                    "last_update": str(dttm.now()),
                                    "fetched": False,
                                    "to_pdf": False,
                                    "pdf_file_path": None
                                }
                            }
                        }
                    }
                }

        # this will give fortnight quiz pdf (218052, Sept 1-15 onwards usable)
        elif product_name.find('Current Affairs [PDF') == 0 and int(product_id) > 218051:
            logging.debug("Product is of type `Current Affairs [PDF`")
            _yr, _mth, _dt = gk_utils.get_year_month_date(product_name.split('-', 1)[1].replace(" ", ""))
            _dt = month_map[_mth.title()] + ". " + _dt
            try:
                if all_files_metadata["current_affairs"]["fortnight"][_yr][_mth][_dt]["product_id"] == product_id:
                    logging.debug("Product has already been fetched")
                    continue
            except KeyError:
                # when some key doesn't exist -> that this is a new product on the page
                logging.info("[New Product Fortnight] | " + str(product_id) + ": " + product_name)
                flag = 0
                temp = {
                    "current_affairs": {
                        "fortnight": {
                            _yr: {
                                _mth: {
                                    _dt: {
                                        "product_id": product_id,
                                        "last_update": str(dttm.now()),
                                        "fetched": False,
                                        "to_pdf": False,
                                        "pdf_file_path": None
                                    }
                                }
                            }
                        }
                    }
                }

        # this will find upsc prelims PDFs (usable after august 2019 - 344039)
        elif product_name.find('UPSC Prelims Current Affairs Quiz [PDF]') == 0:
            logging.debug("Product is of type `UPSC Prelims Current Affairs Quiz [PDF]`")
            _yr, _mth, _dt = gk_utils.get_year_month_date(product_name.split("-")[1].replace(" ", ""))
            _mth = month_map[_mth.title()] + ". " + _mth.title()
            try:
                if all_files_metadata["current_affairs"]["upsc_prelims"][_yr][_mth]["product_id"] == product_id:
                    logging.debug("Product has already been fetched")
                    continue
            except KeyError:
                # when some key doesn't exist -> that this is a new product on the page
                logging.info("[New Product UPSC Prelims] | " + str(product_id) + ": " + product_name)
                flag = 0
                temp = {
                    "upsc_prelims": {
                        "monthly": {
                            _yr: {
                                _mth: {
                                    "product_id": product_id,
                                    "last_update": str(dttm.now()),
                                    "fetched": False,
                                    "to_pdf": False,
                                    "pdf_file_path": None
                                }
                            }
                        }
                    }
                }

        # these are not useful
        else:
            logging.debug("Product is of type `others`")
            product_name = product_name.replace(" ", "")
            try:
                if all_files_metadata["gktoday"][product_name]['id'] == product_id:
                    logging.debug("Product has already been fetched")
                    continue
            except KeyError:
                # when some key doesn't exist -> that this is a new product on the page
                logging.info("[New Product Unusable] | " + str(product_id) + ": " + product_name)
                flag = 0
                temp = {"gktoday": {product_name: {"id": product_id}}}

        insert_into_dict(all_files_metadata, temp)
    return flag


# function to find product details in a gage when using fetch_product_metadat with classic mode
def get_res_classic_fetch(html):
    # parse page using beautiful soup
    soup = BeautifulSoup(html, 'lxml')

    # get all articles from the current page
    table_html = soup.find('main', {"id": "main"}).find('ul')
    all_products = table_html.find_all('li')
    
    return all_products


# extract the metadata from the main_page_json to get the details of all the available products
# this method calls the json API to get the JSON Directly
def fetch_product_metadata(classic=False):
    logging.debug("starting fetch_product_metadata():")
    i = 1
    page_url = main_page_json
    if classic:
        page_url = main_page_url
    flag = 1
    while True:
        fetch_url = page_url.format(page_no=i)
        logging.info("fetching Page:" + str(i))
        res = requests.get(url=fetch_url, headers=req_head, allow_redirects=False)
        # when there are no.
        if res.status_code == 301:
            res = requests.get(url=res.headers['Location'], headers=req_head)
        # when there are no more posts to get
        if res.status_code == 400 or res.status_code == 302:
            log_msg = "No more products to fetch.\nMode: "
            if classic:
                log_msg += "Classic."
            else:
                log_msg += "New."
            log_msg += " Page: " + str(i)
            logging.info(log_msg)
            break
        # load the result into a dict
        if classic:
            try:
                products = get_res_classic_fetch(res.text)
            except AttributeError as err:
                logging.exception("couldn't find the product table (ul) or (li) in the page: " + res.url)
                raise err
        else:
            try:
                products = json.loads(res.text)
            except json.decoder.JSONDecodeError:
                products = json.loads(res.text[4:])
        flag = categorize_product(all_products=products, classic=classic)
        # when page doesn't contain any new product, simply break
        if flag:
            logging.debug("No new Product found on the current page. Skipping all other pages")
            break
        i += 1
    # if some changes have been made to all_files_metadata.json
    # we need to write the file to disk
    if not flag:
        logging.debug("Writing all_files_metadata.json to Disk")
        write_data(quiz_folder + "all_files_metadata.json", json.dumps(all_files_metadata))


# main function
def quiz_run(classic=False):
    try:
        # initialize metadata
        init_metadata()
        _log_level = logging.INFO
        _log_file_path = LOG_PATH + "scraping_quiz.log"

        logger = logging.getLogger()
        logger.setLevel(_log_level)
        file_handler = logging.FileHandler(_log_file_path)
        file_handler.setLevel(_log_level)
        formatter = logging.Formatter('%(levelname)s:%(asctime)s:%(name)s:%(message)s')
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        # logging.basicConfig(filename=LOG_PATH + "scraping_quiz.log", level=logging.INFO)
        fetch_product_metadata(classic)
        fetch_all_products(True)
        logging.info("Scraping successful!")
    except Exception as e:
        logging.exception('Error In __main__')
        raise


if __name__ == "__main__":
    quiz_run()