# daily current url = https://currentaffairs.gktoday.in/wp-json/wp/v2/dailycurrent
from bs4 import BeautifulSoup  # BS for parsing HTML
from datetime import datetime as dttm, time  # datetime for logging
import os
import json
import docx
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.shared import Cm
from scraping.gk_utility import docxToPDF
import logging
from scraping.gk_utility import gk_utils
from scraping.gk_utility.gk_utils import fetch_article_soup
from scraping import main_page, MAGAZINE_LOG_PATH, magazine_folder, pdf_folder, docx_folder, month_map
from scraping.gk_utility import YEARS, MONTHS


# Log file path
LOG_PATH = MAGAZINE_LOG_PATH
LOG_PATH = os.path.expanduser(LOG_PATH)
# TO expand the "~/ path"
magazine_folder = os.path.expanduser(magazine_folder)

# this contains data of each article
daily_articles_metadata = {}
# this is a temporary list of articles which are fetched from the website. this gets dumped into the daily_articles_json
all_articles = {}
# this contains only metadata of each pdf file, like : whether file is latest, name of file, location, etc
all_files_metadata = {}


def make_directories():
    # to create the folder for docx and pdf files if they doesn't exist already
    os.makedirs(magazine_folder + "docx/", exist_ok=True)
    os.makedirs(magazine_folder + "pdf/", exist_ok=True)
    os.makedirs(LOG_PATH, exist_ok=True)
    return 1


def convert_to_pdf(docx_path):
    try:
        docxToPDF.convert_to(magazine_folder + pdf_folder, docx_path)
        name = docx_path.replace(magazine_folder + docx_folder, "")
        logging.info("Successfully converted `" + name + "` to PDF.")
    except Exception:
        logging.exception('Error In convert_to_pdf()')
        raise


# convert new docx files to pdf
def convert_new_docx_to_pdf():
    for _year, _year_val in all_files_metadata.items():
        for _month, _month_val in _year_val.items():
            if not _month_val['to_pdf']:
                docx_file_name = _month_val['file_name'] + '.docx'
                docx_path = magazine_folder + "docx/" + docx_file_name
                # print("converting")   # testing
                convert_to_pdf(docx_path)
                all_files_metadata[_year][_month]['to_pdf'] = True
                all_files_md_path = magazine_folder + "all_files_metadata.json"
                write_data(all_files_md_path, json.dumps(all_files_metadata))


# convert json string keys to ints if possible
def keys_to_int(x):
    res = {}
    for k, v in x.items():
        try:
            res[int(k)] = v
        except ValueError:
            res[k] = v

    return res


# wrapper for get_year_month_date_wrapper
def get_year_month_date_wrapper(date_string):
    yr, mth, dt = gk_utils.get_year_month_date(date_string)
    return int(yr), mth, int(dt)


# footer page number
# https://stackoverflow.com/a/56676220/5954974
def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fld_char1 = create_element('w:fldChar')
    create_attribute(fld_char1, 'w:fldCharType', 'begin')

    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "PAGE"

    fld_char2 = create_element('w:fldChar')
    create_attribute(fld_char2, 'w:fldCharType', 'end')

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


# create front page of the docx file
def front_page(doc, title, updated_on):
    picture_path = './resources/gktoday_frontpage.jpg'
    p = doc.add_paragraph('')
    front = p.add_run('GKTODAY\n\n\n')
    font = front.font
    font.bold = True
    font.size = Pt(16.5)
    font.name = 'Arial'
    p.paragraph_format.alignment = 1
    font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)

    r = p.add_run()
    r.add_text('\n')

    r.add_picture(picture_path, width=Cm(2.78), height=Cm(3.96))

    docx_title = p.add_run('\n' + title + '\n\n\n\n')
    font = docx_title.font
    font.bold = True
    font.size = Pt(16.5)

    p.paragraph_format.alignment = 1
    font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)

    update = p.add_run('Last Updated: ' + updated_on + '\n' * 23)
    font = update.font
    font.bold = True
    font.size = Pt(12)

    p.paragraph_format.alignment = 1
    font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)

    break_para = doc.add_paragraph('')
    advert = break_para.add_run('निःशुल्क सेवा\n@GKTodayQuizBot on Telegram')
    font = advert.font
    font.bold = True
    font.size = Pt(20)

    break_para.paragraph_format.alignment = 1
    font.color.rgb = docx.shared.RGBColor(0xff, 0x99, 0x33)

    break_para.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)


# method to create contents table page for the docx file
# doc is used to write to the docx file
def contents_page(year, month, doc, monthly_articles):
    p = doc.add_paragraph('')
    front = p.add_run('Contents\n')
    font = front.font
    font.bold = True
    font.size = Pt(16.5)

    p.paragraph_format.alignment = 1
    font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0xff)
    p = doc.add_paragraph()
    for date, date_data in monthly_articles.items():
        if date == 'docx_created' or date == 'last_updated':
            continue
        # here first print the (Month date, Year) heading [Bold, size=12]
        title = str(month.capitalize()) + " " + str(date) + ", " + str(year)
        date_heading = p.add_run(title)
        font = date_heading.font
        font.bold = True
        font.size = Pt(12)

        font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)
        # p = doc.add_paragraph('')
        r = p.add_run('\n')
        font = r.font
        font.size = Pt(9)

        for article_num, article_data in date_data.items():
            # here create a entry of each article's heading with a leading \t [size=9] article_data['post_heading']
            r.add_text("\t" + str(article_num) + '. ' + article_data['post']['post_heading'])
            r.add_break(docx.enum.text.WD_BREAK.LINE)

        p.add_run("\n")
    p.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    return 0


# {
#   "date" : {
#           <article1>,
#           <article2>
#       }
# }
#
# <article>
# {
#   "url": 1,
#   "fetched": True/False,
#   "post": {
#           "post_heading": "some heading",
#           "post_first_para": "first para example",
#           "post_body": {
#               "HEADINGS": {
#                   "text": "some paragraph for the heading ",
#                   "SUBHEADINGS": {
#                       "text": "some paragraph for the subheading"
#                   }
#               }
#           }
#       }
# }
# write data to the docx file
def write_article_data_to_word(year_name, month_name, monthly_articles):
    path = magazine_folder + docx_folder
    file_name = month_map[month_name] + ". " + str(month_name) + ", " + str(year_name) + ".docx"
    try:
        os.makedirs(path, exist_ok=True)
        try:
            doc = docx.Document()
            # this is when you want to append
            # doc=docx.Document(path + file_name)
        except docx.opc.exceptions.PackageNotFoundError:  # file doesn't exist
            doc = docx.Document()
        except FileNotFoundError:
            doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(0.81)
            section.bottom_margin = Cm(0.81)
            section.left_margin = Cm(1.31)
            section.right_margin = Cm(1.27)

        font = doc.styles['Normal'].font
        font.name = 'Verdana'

        header_section = doc.sections[0]
        header = header_section.header
        header_section.header_distance = Cm(0.4)

        footer_section = doc.sections[0]
        footer = footer_section.footer
        footer_section.footer_distance = Cm(0.4)
        add_page_number(footer.paragraphs[0].add_run())
        footer.paragraphs[0].paragraph_format.alignment = 2

        header_text = header.paragraphs[0]
        header_text.paragraph_format.alignment = 2
        advert = header_text.add_run("@GKTodayQuizBot on Telegram")
        font = advert.font
        font.color.rgb = docx.shared.RGBColor(0xff, 0x99, 0x33)
        font.bold = True
        last_date = 0
        for date, _ in monthly_articles.items():
            # skip not useful fields
            if date == 'docx_created' or date == 'last_updated':
                continue
            if last_date < date:
                last_date = date
        title_of_magazine = "Current Affairs [PDF] - " + str(month_name.capitalize()) + " 1-" + str(last_date) + ", "\
                            + str(year_name)
        last_updated = monthly_articles['last_updated']
        front_page(doc, title_of_magazine, last_updated)
        if contents_page(year_name, month_name, doc, monthly_articles):
            return 1

        for date, date_articles in monthly_articles.items():
            # skip not useful fields
            if date == 'docx_created' or date == 'last_updated':
                continue
            # this is for the channel name
            p = doc.add_paragraph('')
            advert = p.add_run("@GKTodayQuizBot on Telegram")
            font = advert.font
            font.bold = True
            font.size = Pt(20)

            p.paragraph_format.alignment = 1
            font.color.rgb = docx.shared.RGBColor(0x00, 0xB0, 0x50)
            # this is for the date of the articles
            p = doc.add_paragraph('')
            full_date = str(month_name.capitalize()) + " " + str(date) + ", " + str(year_name)
            date_run = p.add_run(full_date)  # Date of the current Quiz
            font = date_run.font

            font.bold = True
            font.size = Pt(14)
            p.paragraph_format.alignment = 1
            font.color.rgb = docx.shared.RGBColor(0x17, 0x36, 0x5D)
            p.runs[0].add_break(docx.enum.text.WD_BREAK.LINE)
            hr = p.add_run('_' * 77)
            font = hr.font
            font.color.rgb = docx.shared.RGBColor(0x31, 0x84, 0x9B)

            for article_num, article_data in date_articles.items():
                p = doc.add_paragraph('')
                article_data = article_data['post']
                art_heading = p.add_run(str(article_num) + '. ' + str(article_data['post_heading']))
                font = art_heading.font

                font.underline = True
                font.color.rgb = docx.shared.RGBColor(0xC0, 0x00, 0x00)
                p.paragraph_format.alignment = 1
                p.paragraph_format.keep_with_next = True
                font.bold = True
                font.size = Pt(13)

                p = doc.add_paragraph('')
                p.paragraph_format.alignment = 0
                all_text = p.add_run(article_data['post_first_para'])
                font = all_text.font
                font.size = Pt(12)

                for sub_heading, sub_heading_data in article_data['post_body'].items():
                    p = doc.add_paragraph()
                    p.paragraph_format.keep_together = True
                    sub_head = p.add_run(sub_heading)
                    font = sub_head.font
                    font.bold = True

                    font.size = Pt(12)
                    font.color.rgb = docx.shared.RGBColor(0x1F, 0x49, 0x7D)

                    # this case handles the case where subheading contains a text
                    if sub_heading_data['text'] is not None:
                        if type(sub_heading_data['text']) == list:
                            p.paragraph_format.space_after = Pt(1.15)
                            for i in sub_heading_data['text']:
                                p = doc.add_paragraph(i, style="List Bullet")
                                all_text = p.add_run()
                                font = all_text.font
                                font.size = Pt(12)
                        else:
                            p.add_run('\n')
                            all_text = p.add_run(sub_heading_data['text'])
                            font = all_text.font
                            font.size = Pt(12)

                    # sub sub headings
                    for in_sub_head, in_sub_head_data in sub_heading_data.items():
                        if in_sub_head != "text":
                            p = doc.add_paragraph('')
                            p.paragraph_format.keep_together = True
                            sub_head = p.add_run(in_sub_head)
                            font = sub_head.font
                            font.bold = True
                            font.color.rgb = docx.shared.RGBColor(0x94, 0x36, 0x34)
                            font.size = Pt(12)

                            if in_sub_head_data['text'] is not None:
                                if type(in_sub_head_data['text']) == list:
                                    p.paragraph_format.space_after = Pt(1.15)
                                    for i in in_sub_head_data['text']:
                                        p = doc.add_paragraph(i, style="List Bullet")
                                        all_text = p.add_run()
                                        font = all_text.font
                                        font.size = Pt(12)
                                else:
                                    p.add_run('\n')
                                    all_text = p.add_run(in_sub_head_data['text'])
                                    font = all_text.font
                                    font.size = Pt(12)
                p = doc.add_paragraph()
                p.add_run("")

        doc.save(path + file_name)  # save the docx file

    except Exception:
        logging.exception('Error In write_to_docx()')
        raise

    return 0


# write data to the file
def write_data(path, data):
    try:
        os.makedirs(magazine_folder, exist_ok=True)
        f = open(path, "w")
        f.write(data)
        f.close()

    except Exception:
        logging.exception('Error In write_data()')
        raise


# initialize the daily_articles dictionary
def init_metadata():
    try:
        make_directories()
        global daily_articles_metadata
        global all_files_metadata
        try:
            f = open(magazine_folder + "daily_articles_metadata.json", "r")
            daily_articles_metadata = json.loads(f.read(), object_hook=keys_to_int)
            f.close()
        except FileNotFoundError:
            pass

        try:
            g = open(magazine_folder + "all_files_metadata.json", "r")
            all_files_metadata = json.loads(g.read(), object_hook=keys_to_int)
            g.close()
        except FileNotFoundError:
            pass

    except Exception:
        logging.exception('Error In init_metadata()')
        raise


# removes all the html tags from the string
def clean_html(strng):
    return BeautifulSoup(strng, "lxml").text


# format the date properly remove unwanted characters
def fix_date(dt):
    if dt.find('&') != -1:
        dt = dt.replace('&', '-')
    elif dt.find(' &') != -1:
        dt = dt.replace(' &', '-')
    elif dt.find('& ') != -1:
        dt = dt.replace('& ', '-')

    return dt


def get_depth(heading):
    depth = 0
    while len(heading) == 1 and heading.findChildren(recursive=False)[0]\
            .find('a', recursive=False) is None and heading.name != "a":
        depth += 1
        heading = heading.contents[0]
    return depth


def peal(heading):
    count = get_depth(heading)
    while len(heading) == 1 and heading.find('a', recursive=False) is None and count > 0:
        heading = heading.contents[0]
        count -= 1
    return heading


# function to create the nested json object from reading the soup file
def create_data_object(all_sub_headings_toc, post, num=0, inc=1, decimal_places=0):
    data_obj = {}
    for x in all_sub_headings_toc:
        if x.find('a', recursive=False) is not None:
            heading = str(round(num + inc, decimal_places)) + ". " + x.find('a', recursive=False).text
            heading_id = x.find('a', recursive=False)['href']
            heading_id = heading_id.replace("#", "")
            head_text = post.find(id=heading_id).parent.next_sibling.next_sibling
            if head_text.name == "ul":
                temp_list = []
                all_bullets = head_text.find_all('li')
                for y in all_bullets:
                    temp_list.append(y.text)
                head_text = temp_list
            elif head_text.name != "p":
                head_text = None
            else:
                head_text = head_text.text
            to_rem = x.find('a', recursive=False)
            x = x.findChildren(recursive=False)
            x.remove(to_rem)
            if len(x) == 1:
                x = x[0]

        x = peal(x)
        inside_data_obj = {}
        num += inc
        if len(x) >= 1:
            if x.find('a', recursive=False) is not None:
                inside_data_obj = create_data_object(x, post, num, inc * 0.1, decimal_places + 1)
            else:
                inside_data_obj = create_data_object(x.contents, post, num, inc * 0.1, decimal_places + 1)
        inside_data_obj['text'] = head_text
        data_obj[heading] = inside_data_obj

    return data_obj


# EXAMPLE
# {
#   "sub_heading": {
#     "text": "text of subheading",
#     "sub_sub_heading1": {
#       "text": "text of sub sub heading 1"
#     },
#     "sub_sub_heading2": {
#       "text": "text of sub sub heading 2"
#     }
#   }
# }
# fallback method for creating data objects when table of contents isn't available
def create_data_object_fallback(post_soup):
    start = post_soup.find("div", {"class": "featured_image"})
    heading = start.next_sibling.next_sibling.next_sibling.next_sibling.text
    data = start.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling
    data_obj = {}

    if data.name == 'ul':
        all_lis = data.find_all("li")
        data = []
        for i in all_lis:
            data.append(i.text)
    else:
        data = data.text

    data_obj[heading] = {}
    data_obj[heading]['text'] = data

    return data_obj


# function to extract relevant article data from the soup object aotm : article of the month mota : month of the
# article, if mota doesnt match the month in the soup, return -1 (implies that we need to change mth value of the
# fetch function) else return 0 (normal) mth : month for which we are currently creating the magazine yr : year for
# which we are currently creating the
#
# data example:
# {
#   "post_heading": "heading here",
#   "post_first_para": "some small summary here",
#   "post_body": {
#       }
# }
def extract_data_from_soup(soup, mth):
    # implement this make different provisions if the aotm is 1, i.e. first article of the month, set bigger heading
    # for the month, different font color etc
    post = soup.find("div", {"class": "middle_content"})
    heading = post.find("h1").text
    # if contents table isn't available
    try:
        # table of content
        toc = post.find("ul", {"class": "toc_list"})
        all_sub_headings_toc = toc.find_all("li", recursive=False)
        data_obj = create_data_object(all_sub_headings_toc, post)
    except AttributeError:
        logging.info("TOC isn't available for the current post.")
        data_obj = create_data_object_fallback(post)

    date_of_article = post.find("span", {"class": "meta_date"}).text
    text = post.find("p").text
    all_sub_headings = post.find_all("h4")
    all_paras = []
    for sub_heading in all_sub_headings:
        next_p = sub_heading.next_sibling
        while next_p == "\n":
            next_p = next_p.next_sibling
        para = next_p.text + "\n"
        all_paras.append(para)
    sub_headings = []
    for sub_heading in all_sub_headings:
        sub_headings.append(sub_heading.text)

    art_yr, art_mth, art_dt = get_year_month_date_wrapper(date_of_article)
    date_of_article = art_mth + " " + str(art_dt) + ", " + str(art_yr)

    article_data = {"post_heading": heading, "post_first_para": text, "post_body": data_obj}

    # if article is published in another month for which we are collecting the data
    if art_mth.lower() != mth.lower():
        article_data["article_num"] = 1
        return article_data, date_of_article, -1
    return article_data, date_of_article, 0


# function which checks daily_articles_metadata for the last article which was stored in it and returns the url
# returns `None` is url not present (i.e. dict have never been initialized)
def get_last_fetched_url():
    try:
        if "url" in daily_articles_metadata["last_fetched_url"]:
            yr = daily_articles_metadata["last_fetched_url"]['year']
            mth = daily_articles_metadata["last_fetched_url"]['month']
            dt = daily_articles_metadata["last_fetched_url"]['date']
            post_no = daily_articles_metadata["last_fetched_url"]['post_no']
            url = daily_articles_metadata["last_fetched_url"]['url']
            return int(yr), mth, int(dt), int(post_no), url
    except KeyError:
        return None, None, None, None, None
    return None, None, None, None, None


# function to set daily_articles_metadata for the last article
def set_last_fetched_url(yr, mth, dt, post_no, url, fetch_time):
    try:
        if "last_fetched_url" not in daily_articles_metadata:
            daily_articles_metadata["last_fetched_url"] = {}

        daily_articles_metadata["last_fetched_url"]['year'] = yr
        daily_articles_metadata["last_fetched_url"]['month'] = mth
        daily_articles_metadata["last_fetched_url"]['date'] = dt
        daily_articles_metadata["last_fetched_url"]['post_no'] = post_no
        daily_articles_metadata["last_fetched_url"]['url'] = url
        daily_articles_metadata["last_fetched_url"]['fetch_time'] = fetch_time

    except KeyError:
        logging.exception('Error In set_last_fetched_url()')
        raise


# check key type
# return 1 : key error
def check_key(key):
    if type(key) != str and type(key) != int and type(key) is not None:
        print(key)
        return 1
    return 0


# function to write each post to the json file
def write_post_to_json(post_no, date, url=None, post=None):
    yr, mth, dt = get_year_month_date_wrapper(date)
    yr, mth, dt = int(yr), mth, int(dt)
    if check_key(yr):
        print("post_no:", post_no, "date:", date, "url:", url)
    if check_key(mth):
        print("post_no:", post_no, "date:", date, "url:", url)
    if check_key(dt):
        print("post_no:", post_no, "date:", date, "url:", url)
    if check_key(post_no):
        print("post_no:", post_no, "date:", date, "url:", url)
    if check_key(url):
        print("post_no:", post_no, "date:", date, "url:", url)

    if yr not in daily_articles_metadata:
        daily_articles_metadata[yr] = {}
    if mth not in daily_articles_metadata[yr]:
        daily_articles_metadata[yr][mth] = {}
    if dt not in daily_articles_metadata[yr][mth]:
        daily_articles_metadata[yr][mth][dt] = {}
    if post_no not in daily_articles_metadata[yr][mth][dt]:
        daily_articles_metadata[yr][mth][dt][post_no] = {}
    if url is not None:
        daily_articles_metadata[yr][mth][dt][post_no]['url'] = url
    if post is not None:
        daily_articles_metadata[yr][mth][dt][post_no]['fetched'] = True
        daily_articles_metadata[yr][mth][dt][post_no]['post'] = post
    else:
        daily_articles_metadata[yr][mth][dt][post_no]['fetched'] = False
        daily_articles_metadata[yr][mth][dt][post_no]['post'] = None

    # this tells the other methods whether this month has been modified after creating a docx file.
    daily_articles_metadata[yr][mth]['docx_created'] = False
    daily_articles_metadata[yr][mth]['last_updated'] = dttm.now().strftime("%B %d, %Y")

    if url is None:
        url = daily_articles_metadata[yr][mth][dt][post_no]['url']
    fetch_time = dttm.now().strftime("%d/%m/%Y %H:%M:%S")
    set_last_fetched_url(yr, mth, dt, post_no, url, fetch_time)


# function which fetches the initial article for a given month and returns it's url
def get_start_url(_yr=1, _mth=0):
    if gk_utils.MONTHS[_mth] == "January":
        _yr -= 1
    _mth -= 1
    mth = gk_utils.MONTHS[_mth]
    yr = YEARS[_yr]
    main_page_url = main_page.format(url_month=mth, url_year=yr)
    _mth += 1
    if _mth == 0:
        _yr += 1

    soup = fetch_article_soup(main_page_url)
    last_article_soup = soup.find('div', class_='post-content').find('h2').find('a')['href']
    soup = fetch_article_soup(last_article_soup)

    try:
        dt = soup.find('span', {"class": "meta_date"}).text
    except AttributeError:
        dt = None
    yr, mth, dt = get_year_month_date_wrapper(dt)
    # get the url of the first article of the month
    # put inside a try except statement (current article may be the last article of the month)
    try:
        next_article_url = soup.find('a', string='Next Update')['href']
    except IndexError:
        next_article_url = None
    return yr, mth, dt, 1, next_article_url


# fetch and store article urls
# force : force the function to use the arguments provided for year and month instead of using the old saved urls
# return 0 means successful
def fetch_articles(_yr=1, _mth=0, force=False):
    yr, mth, dt, post_no, start_url = get_last_fetched_url()
    if start_url is not None:
        _yr, _mth = gk_utils.map_year_month(yr, mth)
        temp_soup = fetch_article_soup(start_url)
        date_of_article = temp_soup.find("span", {"class": "meta_date"}).text

        yr, mth, dt = get_year_month_date_wrapper(date_of_article)
        try:
            next_url = temp_soup.find('a', string='Next Update')['href']
        except TypeError:
            logging.info("no more urls to scrape")
            return 0
        start_url = temp_soup.find('a', string='Next Update')['href']
        mth, post_no, _ = gk_utils.get_month_post_no(next_url, mth, post_no)

    if start_url is None or force:
        yr, mth, dt, post_no, start_url = get_start_url(_yr, _mth)

        if start_url is not None:
            mth, post_no, temp_soup = gk_utils.get_month_post_no(start_url, mth, post_no)
            start_url = temp_soup.find('a', string='Next Update')['href']

    # cases when the year, month has no article present
    if start_url is None:
        return 1

    # article of the month
    aotm = post_no
    current_article_url = start_url

    while current_article_url:
        article_soup = fetch_article_soup(current_article_url)
        art_data, art_date, res = extract_data_from_soup(article_soup, mth)

        if article_soup is not None:
            aotm += 1
        else:
            # there might have been some issue with the
            time.sleep(5)
            continue

        # res = -1 means that: either no more posts are there, or month have changed
        if res == -1:
            aotm = 2
            _mth += 1
            _mth %= 11
            if _mth == 0:
                _yr += 1
            mth = MONTHS[_mth]

        logging.info(str(aotm - 1) + " " + str(mth) + " " + ": " + current_article_url)
        print(aotm - 1, mth, ": ", current_article_url)

        # this will write the metadata of the article to the dict `daily_articles_metadata`
        write_post_to_json(aotm - 1, art_date, current_article_url, art_data)
        save_path = str(magazine_folder) + "daily_articles_metadata.json"
        write_data(path=save_path, data=json.dumps(daily_articles_metadata))

        # testing
        try:
            current_article_url = article_soup.find('a', string='Next Update')['href']
        except IndexError:
            current_article_url = None
        except TypeError:
            print("[" + str(dttm.now()) + "] no more urls to scrape!")
            return 0

    return 0


# write to all_files_metadata
def write_all_files_metadata(year, month, file_name, file_path, update_time):
    month = month
    if year not in all_files_metadata:
        all_files_metadata[year] = {}
    if month not in all_files_metadata[year]:
        all_files_metadata[year][month] = {}
    all_files_metadata[year][month]['file_name'] = file_name
    all_files_metadata[year][month]['file_path'] = file_path
    all_files_metadata[year][month]['last_updated'] = update_time
    all_files_metadata[year][month]['to_pdf'] = False


# this will create docx file for each of the months using the `daily_articles_metadata` object
def create_docx(force=False):
    # year can also contain `last_fetch_url`

    for year, year_data in daily_articles_metadata.items():
        if year != 'last_fetched_url':
            for month, month_data in year_data.items():
                file_name = month_map[month] + ". " + str(month) + ", " + str(year)
                pdf_path = str(magazine_folder) + "pdf/" + file_name
                docx_path = str(magazine_folder) + "docx/" + file_name
                if month_data['docx_created'] is False or not os.path.exists(docx_path + ".docx") or force:
                    # do writing stuff now, then change `docx_created` to `True`
                    if write_article_data_to_word(year, month, month_data):
                        return 1  # error
                    daily_articles_metadata[year][month]['docx_created'] = True
                    update_time = dttm.now().strftime("%d/%B/%Y %H:%M:%S")
                    write_all_files_metadata(int(year), month, file_name, pdf_path + ".pdf", update_time)
    save_path = str(magazine_folder) + "daily_articles_metadata.json"
    write_data(path=save_path, data=json.dumps(daily_articles_metadata))
    write_data(path=str(magazine_folder) + "all_files_metadata.json", data=json.dumps(all_files_metadata))
    return 0  # success


# main function
def magazine_run():
    try:
        # initialize metadata
        init_metadata()

        _log_level = logging.INFO
        _log_file_path = LOG_PATH + "scraping_magazine.log"

        logger = logging.getLogger()
        logger.setLevel(_log_level)
        file_handler = logging.FileHandler(_log_file_path)
        file_handler.setLevel(_log_level)
        formatter = logging.Formatter('%(levelname)s:%(asctime)s:%(name)s:%(message)s')
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

        # also provide year and month if you wanna start from there
        if fetch_articles():
            logging.error("There was some error in fetching articles")
        else:
            logging.info("Articles fetch successful")

        # create doc files
        if create_docx(force=False):
            logging.error("There was some error in creating docx files")
        else:
            logging.info("Docx Creation Successful")
        # Convert new docx files to pdf
        convert_new_docx_to_pdf()
        logging.info("Scraping Successful!")

    except Exception as err:
        logging.exception('Error In magazine_run()')
        raise err


if __name__ == "__main__":
    magazine_run()
